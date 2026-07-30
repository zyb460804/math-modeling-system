#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""保留 Word 格式替换段落（AIGC 降重专用）。

融合自 huanghfzhufeng/aigc-deslop（18★）的 replace_docx.py。
场景：aigc-reduce 把论文段落改写降 AIGC 后，需要把改写后的文字放回原 .docx，
      但不能破坏原排版（字体/字号/行距/页边距）。本脚本按章节标题定位，
      章节内段落按顺序对应 markdown，原地替换保留 run 格式。

实测：9 轮 PaperPure 迭代，AIGC 率 55% → 11%。

用法：
  python replace_docx_preserve_format.py 原.docx 改后.md 输出.docx
"""
import re
import sys
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def normalize_title(text: str) -> str:
    """标准化标题用于匹配：去所有空白、保留中英数字符号"""
    return re.sub(r"\s+", "", text.strip())


def parse_md_sections(md_path: str) -> dict:
    """解析 markdown，按章节切段。返回 {标题(标准化): [段落]}"""
    sections = {}
    current_title = None
    current_paras = []

    with open(md_path, "r", encoding="utf-8") as f:
        for line in f:
            s = line.rstrip("\n")
            stripped = s.strip()

            # 标题
            m = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if m:
                if current_title is not None:
                    sections[current_title] = current_paras
                current_title = normalize_title(m.group(2))
                current_paras = []
                continue

            # 跳过空行
            if not stripped:
                continue
            # 跳过表格、列表、引用、图表注、关键词
            if stripped.startswith(("|", "-", "*", ">")):
                continue
            if re.match(r"^\d+\.\s", stripped):
                continue
            if stripped.startswith(("表", "图")) and len(stripped) < 30:
                continue
            if stripped.startswith("**"):
                continue
            if len(stripped) < 30:
                continue

            current_paras.append(stripped)

        if current_title is not None:
            sections[current_title] = current_paras

    return sections


_REF_PATTERN = re.compile(r"(\[\d+\])")  # 每个 [N] 单独成 run，匹配原文 docx 习惯


def _add_run_with_font(para, text, font_name, font_size, superscript=False):
    """添加一个 run，应用字体和上标设置"""
    run = para.add_run(text)
    if font_name:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size
    if superscript:
        run.font.superscript = True
    return run


def replace_paragraph_text(para, new_text: str):
    """替换段落文本，保留段落级格式。
    特殊处理：
    - \\n 软换行
    - [N] 或 [N][M] 自动识别为上标（学术引用编号）
    """
    # 保存第一个 run 的关键格式
    font_name = font_size = None
    if para.runs:
        first = para.runs[0]
        font_name = first.font.name
        font_size = first.font.size

    # 删除所有 runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    # 按软换行分行处理
    lines = new_text.split("\n")
    last_run = None
    for line_idx, line in enumerate(lines):
        if line_idx > 0 and last_run is not None:
            last_run.add_break()

        # 在每行内，按 [N] 拆分，引用编号设为上标
        parts = _REF_PATTERN.split(line)
        for part in parts:
            if not part:
                continue
            is_ref = bool(_REF_PATTERN.fullmatch(part))
            last_run = _add_run_with_font(
                para, part, font_name, font_size, superscript=is_ref
            )


def should_skip(para) -> tuple:
    """
    判断段落是否应跳过不替换。
    返回 (skip: bool, reason: str)
    """
    text = para.text.strip()

    # 1. 太短跳过（图表注、空段）
    if len(text) < 30:
        return True, "太短"

    # 2. 段内有真实多格式差异 → 保留原结构（关键词、Keywords 等）
    sigs = set()
    for run in para.runs:
        rpr = run._element.find(qn("w:rPr"))
        ea = None
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                ea = rfonts.get(qn("w:eastAsia"))
        sigs.add((
            run.font.name, ea, run.font.size,
            run.font.bold, run.font.italic, run.font.underline,
        ))
    if len(sigs) > 1:
        return True, "段内多格式（保留原结构）"

    # 3. 以特定关键字开头跳过
    skip_prefixes = (
        "关键词", "Keywords", "关键字", "ABSTRACT",
        "Based on", "Design and", "图", "表",
    )
    for prefix in skip_prefixes:
        if text.startswith(prefix):
            return True, f"以 '{prefix}' 开头"

    # 4. 英文段跳过（中文比例 < 40% 视为英文）
    chinese_count = sum(1 for c in text if "一" <= c <= "鿿")
    if chinese_count / len(text) < 0.4:
        return True, "英文段（中文 < 40%）"

    # 5. 大量数字段跳过（< 30% 中文 + 大量数字 = 表格内容或公式）
    digit_count = sum(1 for c in text if c.isdigit())
    if digit_count > len(text) * 0.3 and chinese_count < len(text) * 0.3:
        return True, "数字主导段"

    return False, ""


def count_section_replaceable(doc):
    """第一遍扫描：统计每个章节内"可替换的 Normal 长段"数"""
    counts = {}
    current = None
    for para in doc.paragraphs:
        style = para.style.name
        if "toc" in style.lower():
            continue
        if "Heading" in style:
            current = normalize_title(para.text.strip())
            counts.setdefault(current, 0)
            continue
        if style != "Normal" or current is None:
            continue
        skip, _ = should_skip(para)
        if not skip:
            counts[current] = counts.get(current, 0) + 1
    return counts


def fit_md_to_docx(md_paras, docx_capacity):
    """把 markdown 段调整为正好 docx_capacity 段：多余的合并到最后一段"""
    if docx_capacity <= 0:
        return []
    if len(md_paras) <= docx_capacity:
        return md_paras
    # 把超出的段合并到最后（用换行连接，python-docx 会处理为段内软换行）
    fitted = list(md_paras[: docx_capacity - 1])
    fitted.append("\n".join(md_paras[docx_capacity - 1 :]))
    return fitted


def replace_docx(src_docx: str, md_path: str, dst_docx: str):
    sections = parse_md_sections(md_path)

    print(f"markdown 章节数：{len(sections)}")
    for t, paras in sections.items():
        print(f"  [{t}] {len(paras)} 段")

    doc = Document(src_docx)

    # 第一遍：统计每章节 docx 可替换段数
    docx_capacity = count_section_replaceable(doc)
    print()
    print("docx 章节可替换段数：")
    for title, cnt in docx_capacity.items():
        if cnt > 0:
            md_count = len(sections.get(title, []))
            flag = "✓" if md_count == cnt else "⚙ 自动合并" if md_count > cnt else "⚠ md 不够"
            print(f"  [{title}] docx:{cnt} md:{md_count} {flag}")

    current_section_paras = None
    para_idx = 0
    replaced = 0
    skipped_chapters = []
    section_stats = []  # (章节, docx段数, md段数, 是否一致)

    current_section_title = None
    current_section_replaced = 0
    current_section_md_total = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # toc 跳过
        if "toc" in style.lower():
            continue

        # Heading：切换章节
        if "Heading" in style:
            # 提交前一章节统计
            if current_section_title is not None:
                section_stats.append((
                    current_section_title,
                    current_section_replaced,
                    current_section_md_total,
                ))

            title = normalize_title(text)
            if title and title in sections:
                # 把 markdown 段调整为 docx 能容纳的段数
                capacity = docx_capacity.get(title, 0)
                current_section_paras = fit_md_to_docx(sections[title], capacity)
                para_idx = 0
                current_section_title = title
                current_section_md_total = len(current_section_paras)
                current_section_replaced = 0
            else:
                current_section_paras = None
                current_section_title = None
                if title:
                    skipped_chapters.append(title)
            continue

        # Caption / 其它非正文
        if style == "Caption":
            continue
        if style != "Normal":
            continue

        # 应用保护规则
        skip, reason = should_skip(para)
        if skip:
            continue

        # 替换
        if current_section_paras and para_idx < len(current_section_paras):
            new_text = current_section_paras[para_idx]
            replace_paragraph_text(para, new_text)
            replaced += 1
            para_idx += 1
            current_section_replaced += 1

    # 提交最后一章
    if current_section_title is not None:
        section_stats.append((
            current_section_title,
            current_section_replaced,
            current_section_md_total,
        ))

    doc.save(dst_docx)

    print()
    print(f"=== 替换完成：{replaced} 段 ===")
    print()
    print("章节替换详情：")
    for title, did, total in section_stats:
        flag = "✓" if did == total else "⚠"
        print(f"  {flag} [{title}] 替换 {did}/{total}")

    if skipped_chapters:
        print()
        print(f"跳过的章节（markdown 中未找到匹配）：")
        for t in skipped_chapters[:10]:
            print(f"  - {t}")

    print()
    print(f"输出文件：{dst_docx}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print(__doc__)
        sys.exit(1)
    replace_docx(sys.argv[1], sys.argv[2], sys.argv[3])
