"""自动化论文评分脚本：对照评分量表对论文进行结构化评分。

用法：python auto_score.py --paper paper_output/final_paper.docx
      python auto_score.py --paper paper_output/final_paper.docx --results-dir paper_output/results

评分标准：7 模块 100 分制（来自 outputs/scoring_rubric.md）
输出：paper_output/results/score_report.json + score_report.md
"""
import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None


# ===== 评分量表（7 模块 100 分制）=====
RUBRIC = {
    "审题": {
        "max": 15,
        "checks": [
            {"name": "题目目标识别", "max": 5, "keywords": ["目标", "目的", "问题", "要求"]},
            {"name": "题型判断", "max": 5, "keywords": ["评价", "预测", "优化", "分类", "聚类", "图论", "仿真"]},
            {"name": "约束与隐含条件", "max": 5, "keywords": ["约束", "条件", "假设", "限制", "指标"]},
        ],
    },
    "建模": {
        "max": 20,
        "checks": [
            {"name": "问题抽象", "max": 5, "keywords": ["模型", "数学", "公式", "函数", "变量"]},
            {"name": "模型假设", "max": 5, "keywords": ["假设", "简化", "不计", "忽略", "考虑"]},
            {"name": "模型结构", "max": 5, "keywords": ["目标函数", "约束条件", "决策变量", "优化"]},
            {"name": "模型链条", "max": 5, "keywords": ["递进", "关联", "耦合", "多阶段"]},
        ],
    },
    "求解与算法": {
        "max": 20,
        "checks": [
            {"name": "方法匹配度", "max": 8, "keywords": ["线性规划", "整数规划", "遗传算法", "粒子群", "TOPSIS", "AHP", "ARIMA"]},
            {"name": "推导完整性", "max": 6, "keywords": ["推导", "证明", "公式", "步骤", "过程"]},
            {"name": "可复现性", "max": 6, "keywords": ["代码", "算法", "实现", "运行", "结果"]},
        ],
    },
    "结果分析": {
        "max": 15,
        "checks": [
            {"name": "结果可信度", "max": 5, "keywords": ["验证", "检验", "合理", "符合", "一致"]},
            {"name": "现实解释", "max": 5, "keywords": ["实际", "意义", "解释", "应用", "价值"]},
            {"name": "结论支撑", "max": 5, "keywords": ["结论", "总结", "建议", "推广"]},
        ],
    },
    "稳健性/误差分析": {
        "max": 10,
        "checks": [
            {"name": "灵敏度分析", "max": 5, "keywords": ["灵敏度", "敏感", "参数变化", "影响"]},
            {"name": "误差说明", "max": 5, "keywords": ["误差", "局限", "不足", "改进"]},
        ],
    },
    "论文写作": {
        "max": 15,
        "checks": [
            {"name": "摘要质量", "max": 5, "keywords": ["摘要", "本文", "针对", "建立", "结果"]},
            {"name": "结构清晰度", "max": 5, "keywords": ["一、", "二、", "三、", "四、", "五、"]},
            {"name": "语言风格", "max": 5, "keywords": ["参考文献", "公式", "图表"]},
        ],
    },
    "答辩": {
        "max": 5,
        "checks": [
            {"name": "问题理解", "max": 2, "keywords": ["分析", "理解", "解释"]},
            {"name": "方法选择", "max": 2, "keywords": ["选择", "原因", "优势"]},
            {"name": "风险意识", "max": 1, "keywords": ["风险", "局限", "改进"]},
        ],
    },
}


def extract_text_from_docx(path: str) -> str:
    """从 docx 提取全文。"""
    if Document is None:
        return ""
    try:
        doc = Document(path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            texts.append(para.text)
        return "\n".join(texts)
    except Exception:
        return ""


def extract_text_from_md(path: str) -> str:
    """从 markdown 提取全文。"""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""


def count_figures(text: str) -> int:
    """统计图表数量。"""
    fig_patterns = [
        r"图\s*\d+", r"Fig\s*\d+", r"Figure\s*\d+",
        r"!\[", r"\.png", r"\.jpg",
    ]
    count = 0
    for pat in fig_patterns:
        count += len(re.findall(pat, text, re.IGNORECASE))
    return count // 2  # 每个图通常出现 2 次（引用+图本身）


def count_tables(text: str) -> int:
    """统计表格数量。"""
    table_patterns = [
        r"表\s*\d+", r"Table\s*\d+",
    ]
    count = 0
    for pat in table_patterns:
        count += len(re.findall(pat, text, re.IGNORECASE))
    return count // 2


def count_references(text: str) -> int:
    """统计参考文献数量。"""
    return len(re.findall(r"\[\d+\]", text))


def count_formulas(text: str) -> int:
    """统计公式数量。"""
    return len(re.findall(r"\$.*?\$", text))


def score_module(module_name: str, module_spec: dict, text: str) -> dict:
    """对单个模块评分。"""
    checks = module_spec["checks"]
    total_score = 0
    total_max = 0
    check_results = []

    for check in checks:
        check_name = check["name"]
        check_max = check["max"]
        keywords = check["keywords"]

        # 统计关键词出现次数
        matches = 0
        for kw in keywords:
            matches += len(re.findall(re.escape(kw), text, re.IGNORECASE))

        # 按匹配度评分
        if matches >= 5:
            score = check_max
        elif matches >= 3:
            score = check_max * 0.8
        elif matches >= 1:
            score = check_max * 0.5
        else:
            score = 0

        score = round(score, 1)
        total_score += score
        total_max += check_max

        check_results.append({
            "name": check_name,
            "max": check_max,
            "score": score,
            "matches": matches,
        })

    return {
        "module": module_name,
        "max": module_spec["max"],
        "score": round(total_score, 1),
        "checks": check_results,
    }


def score_paper(paper_path: str, results_dir: str = None) -> dict:
    """对论文进行完整评分。"""
    # 提取文本
    if paper_path.endswith(".docx"):
        text = extract_text_from_docx(paper_path)
    elif paper_path.endswith(".md"):
        text = extract_text_from_md(paper_path)
    else:
        return {"error": f"Unsupported file format: {paper_path}"}

    if not text:
        return {"error": "Failed to extract text from paper"}

    # 基础统计
    char_count = len(re.sub(r"\s+", "", text))
    figure_count = count_figures(text)
    table_count = count_tables(text)
    reference_count = count_references(text)
    formula_count = count_formulas(text)

    # 逐模块评分
    module_scores = []
    total_score = 0
    for module_name, module_spec in RUBRIC.items():
        result = score_module(module_name, module_spec, text)
        module_scores.append(result)
        total_score += result["score"]

    total_score = round(total_score, 1)

    # 等级判定
    if total_score >= 85:
        grade = "强终稿/冲奖稿"
    elif total_score >= 70:
        grade = "可提交稿"
    elif total_score >= 55:
        grade = "风险稿"
    else:
        grade = "高风险稿"

    # 检查格式要求
    format_issues = []
    if char_count < 18000:
        format_issues.append(f"字数不足：{char_count} < 18000")
    if figure_count < 5:
        format_issues.append(f"图表偏少：{figure_count} < 5")
    if table_count < 3:
        format_issues.append(f"表格偏少：{table_count} < 3")
    if reference_count < 5:
        format_issues.append(f"参考文献偏少：{reference_count} < 5")

    # 构建报告
    report = {
        "generated_at": datetime.now().isoformat(),
        "paper_path": paper_path,
        "total_score": total_score,
        "max_score": 100,
        "grade": grade,
        "statistics": {
            "char_count": char_count,
            "figure_count": figure_count,
            "table_count": table_count,
            "reference_count": reference_count,
            "formula_count": formula_count,
        },
        "module_scores": module_scores,
        "format_issues": format_issues,
    }

    # 加载结果数据（如果提供）
    if results_dir:
        validation_path = os.path.join(results_dir, "validation_report.json")
        if os.path.exists(validation_path):
            with open(validation_path, "r", encoding="utf-8") as f:
                validation = json.load(f)
                report["validation_status"] = validation.get("status", "UNKNOWN")

    return report


def generate_markdown_report(report: dict) -> str:
    """生成 Markdown 格式的评审报告。"""
    lines = [
        "# 论文自动评审报告",
        "",
        f"- 生成时间：{report['generated_at']}",
        f"- 论文路径：{report['paper_path']}",
        "",
        "## 总分",
        "",
        f"**{report['total_score']}/100** — **{report['grade']}**",
        "",
        "## 基础统计",
        "",
        f"- 字数：{report['statistics']['char_count']}",
        f"- 图表：{report['statistics']['figure_count']}",
        f"- 表格：{report['statistics']['table_count']}",
        f"- 参考文献：{report['statistics']['reference_count']}",
        f"- 公式：{report['statistics']['formula_count']}",
        "",
        "## 分项评分",
        "",
        "| 模块 | 满分 | 得分 | 得分率 |",
        "|------|------|------|--------|",
    ]

    for m in report["module_scores"]:
        rate = f"{m['score']/m['max']*100:.0f}%" if m["max"] > 0 else "N/A"
        lines.append(f"| {m['module']} | {m['max']} | {m['score']} | {rate} |")

    lines.append(f"| **合计** | **100** | **{report['total_score']}** | **{report['total_score']}%** |")
    lines.append("")

    # 详细检查
    lines.append("## 详细检查")
    lines.append("")
    for m in report["module_scores"]:
        lines.append(f"### {m['module']}（{m['score']}/{m['max']}）")
        lines.append("")
        for c in m["checks"]:
            status = "✅" if c["score"] >= c["max"] * 0.8 else "⚠️" if c["score"] >= c["max"] * 0.5 else "❌"
            lines.append(f"- {status} {c['name']}：{c['score']}/{c['max']}（匹配 {c['matches']} 次）")
        lines.append("")

    # 格式问题
    if report["format_issues"]:
        lines.append("## 格式问题")
        lines.append("")
        for issue in report["format_issues"]:
            lines.append(f"- ❌ {issue}")
        lines.append("")

    # 验证状态
    if "validation_status" in report:
        lines.append("## 结果验证")
        lines.append("")
        lines.append(f"- 状态：{report['validation_status']}")
        lines.append("")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="Auto paper scoring")
    parser.add_argument("--paper", required=True, help="Paper file path (.docx or .md)")
    parser.add_argument("--results-dir", default=None, help="Results directory for validation")
    parser.add_argument("--output-dir", default="paper_output/results", help="Output directory")
    args = parser.parse_args()

    if not os.path.exists(args.paper):
        print(f"Error: Paper file not found: {args.paper}")
        sys.exit(1)

    # 评分
    report = score_paper(args.paper, args.results_dir)

    if "error" in report:
        print(f"Error: {report['error']}")
        sys.exit(1)

    # 保存 JSON
    os.makedirs(args.output_dir, exist_ok=True)
    json_path = os.path.join(args.output_dir, "score_report.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    # 保存 Markdown
    md_path = os.path.join(args.output_dir, "score_report.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(generate_markdown_report(report))

    # 输出摘要
    print(f"Score: {report['total_score']}/100 — {report['grade']}")
    print(f"  Figures: {report['statistics']['figure_count']}")
    print(f"  Tables: {report['statistics']['table_count']}")
    print(f"  References: {report['statistics']['reference_count']}")
    if report["format_issues"]:
        print(f"  Format issues: {len(report['format_issues'])}")
    print(f"Report: {json_path}")

    return report


if __name__ == "__main__":
    main()
