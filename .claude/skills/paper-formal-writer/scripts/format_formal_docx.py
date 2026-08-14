import argparse
import csv
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path.cwd()
OUTPUT_DIR = BASE_DIR / "paper_output"
SOURCE_FILE = OUTPUT_DIR / "final_paper_source.md"
FALLBACK_SOURCE_FILE = OUTPUT_DIR / "final_paper.md"
OUTLINE_FILE = OUTPUT_DIR / "plan" / "paper_outline.json"
FIGURE_INDEX_FILE = OUTPUT_DIR / "figure_index.json"
TABLE_INDEX_FILE = OUTPUT_DIR / "tables" / "table_index.json"
EVIDENCE_GATE_REPORT = OUTPUT_DIR / "qa" / "evidence_gate_report.json"
DOCX_FILE_FORMAL = OUTPUT_DIR / "final_paper.docx"
DOCX_FILE_DRAFT = OUTPUT_DIR / "final_paper_draft.docx"
REPORT_MD_FORMAL = OUTPUT_DIR / "format_check_report.md"
REPORT_MD_DRAFT = OUTPUT_DIR / "format_draft_report.md"
DOCX_FILE = DOCX_FILE_FORMAL
REPORT_MD = REPORT_MD_FORMAL


# ── 跨 skill 复用 docx-editor-cn 的 LaTeX→OMML（pandoc 链路）─────────────────
# 依赖：pandoc ≥ 2.0（v4.5 体检已装机）。失败时退化为 Cambria Math 纯文本，
# 不阻断生成（但 Word 双击公式不会进编辑器）。
def _locate_docx_editor_scripts() -> Path:
    """CR-7 修复：从本脚本文件位置逐级上溯定位 docx-editor-cn/scripts，
    不再依赖 Path.cwd()——从任意工作目录运行都能 import formula 模块。"""
    here = Path(__file__).resolve()
    for ancestor in here.parents:
        candidate = ancestor / ".claude" / "skills" / "docx-editor-cn" / "scripts"
        if candidate.is_dir():
            return candidate
    # 已知布局兜底（scripts/ → paper-formal-writer/ → skills/）：不存在时
    # import 失败 → _FORMULA_AVAILABLE=False → 报告 DEGRADED（可见，不再静默）。
    return here.parents[2] / "docx-editor-cn" / "scripts"


_DOCX_EDITOR_SCRIPTS = _locate_docx_editor_scripts()
if str(_DOCX_EDITOR_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_DOCX_EDITOR_SCRIPTS))
try:
    from formula import latex_to_omml as _latex_to_omml_pandoc  # type: ignore[import-not-found]
    _FORMULA_AVAILABLE = True
except Exception as _exc:  # pragma: no cover - 环境问题兜底
    _FORMULA_AVAILABLE = False
    print(
        f"[formula] 警告：无法加载 latex_to_omml（{type(_exc).__name__}: {_exc}）。"
        "公式将以 Cambria Math 纯文本呈现，非 Word 原生 OMML。",
        file=sys.stderr,
    )

_OMML_CACHE: dict[str, str] = {}
INLINE_FORMULA_PAT = re.compile(r"\$([^$\n]+?)\$")

# CR-7 修复：OMML 渲染计数（此前零上报，pandoc 缺失=全退化仍报 GENERATED）。
# _append_omml 每次尝试 total+1；成功 ok+1；退化为 Cambria Math 纯文本 fallback+1。
_OMML_STATS: dict[str, int] = {"total": 0, "ok": 0, "fallback": 0}
# H-1/H-8 修复：渲染期失败清单（路径拒绝 / 索引未命中 / 文件缺失），进报告 Failures 节。
_RENDER_FAILURES: list[str] = []


def _latex_to_omml_cached(latex: str) -> str:
    """LaTeX → OMML XML 片段，带缓存。失败返回空串。"""
    if latex in _OMML_CACHE:
        return _OMML_CACHE[latex]
    if not _FORMULA_AVAILABLE:
        _OMML_CACHE[latex] = ""
        return ""
    try:
        omml = _latex_to_omml_pandoc(latex)
        _OMML_CACHE[latex] = omml
        return omml
    except Exception as exc:
        print(f"[formula] LaTeX→OMML 失败 ({latex!r}): {exc}", file=sys.stderr)
        _OMML_CACHE[latex] = ""
        return ""


def _append_omml(paragraph, latex: str) -> bool:
    """把 LaTeX 转 OMML 并 append 到段落元素。成功返回 True。"""
    _OMML_STATS["total"] += 1
    omml_xml = _latex_to_omml_cached(latex)
    if not omml_xml:
        _OMML_STATS["fallback"] += 1
        return False
    try:
        omath = parse_xml(omml_xml)
        paragraph._element.append(omath)
        _OMML_STATS["ok"] += 1
        return True
    except Exception as exc:
        print(f"[formula] OMML 注入失败 ({latex!r}): {exc}", file=sys.stderr)
        _OMML_STATS["fallback"] += 1
        return False


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def load_json(path: Path) -> Any:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def rel(path: Path) -> str:
    try:
        return path.relative_to(BASE_DIR).as_posix()
    except ValueError:
        return path.as_posix()


# H-1 修复：消费点扩展名白名单。图片只收位图/矢量图，表格只收数据文件——
# 防止 md/table_index 一句 ![](任意路径) 把本地文件嵌入外发论文（数据外泄信道）。
IMAGE_EXT_WHITELIST = frozenset({".png", ".jpg", ".jpeg", ".svg", ".gif", ".emf"})
TABLE_EXT_WHITELIST = frozenset({".csv", ".xlsx", ".json"})


def resolve_path(path_text: str, allowed_exts: frozenset[str] | None = None) -> Path:
    """源稿路径 → 实际文件路径。安全门禁（H-1）：

    1. 拒绝绝对路径与含 ``..`` 段的相对路径（Path.parts 检查）；
    2. resolve() 后必须位于 paper_output/ 或项目根内，越界拒绝；
    3. 传入 allowed_exts 时扩展名必须在白名单内。
    任一不满足 → 记入 _RENDER_FAILURES 并返回一个**不存在**的占位路径，
    由下游 add_image/add_table_from_rows 走既有"文件未找到"可见占位，
    不会静默消失，也不会读到白名单外的文件。

    相对路径解析保持原语义：优先 paper_output/（图表实际位置），回退项目根。
    """
    raw = path_text.strip().strip("<>").strip()

    def _rejected(reason: str) -> Path:
        _RENDER_FAILURES.append(f"路径被安全门禁拒绝（{reason}）：{raw}")
        # 只保留文件名部分，保证返回的占位路径本身不可能命中真实文件
        safe_name = Path(raw.replace("\\", "/")).name or "REJECTED"
        return OUTPUT_DIR / "_path_rejected" / safe_name

    if not raw:
        return _rejected("路径为空")
    path = Path(raw)
    if path.is_absolute():
        return _rejected("绝对路径")
    if ".." in path.parts:
        return _rejected("含 .. 上跳段")
    if allowed_exts is not None and path.suffix.lower() not in allowed_exts:
        return _rejected(f"扩展名 {path.suffix or '(无)'} 不在白名单 {sorted(allowed_exts)}")
    for base in (OUTPUT_DIR, BASE_DIR):
        candidate = base / path
        if candidate.exists():
            resolved = candidate.resolve()
            if not (resolved.is_relative_to(OUTPUT_DIR.resolve()) or resolved.is_relative_to(BASE_DIR.resolve())):
                return _rejected("resolve 后位于 paper_output/ 与项目根之外（越界）")
            return resolved
    return OUTPUT_DIR / path  # 文件不存在 → 下游可见占位（既有行为）


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D9D9D9", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margin(cell, margin_twips: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def apply_run_font(run, font_name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)

    for name, font_size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    return text.strip()


# H-10 修复：货币/转义美元误切保护。
# 行内公式切分（INLINE_FORMULA_PAT）之前，先把 `$` 的非公式来源换成不可冲突占位符：
#   1. code span `...`（如 `plot_$x$.py`）——整段原样保护，还原后再由
#      clean_inline_markdown 去反引号；
#   2. `\$` 转义（如 "成本为 \$100"）——按 Markdown 语义还原为字面 `$`。
# 占位符用 Unicode 私有区字符包裹（正文/公式不可能出现），且不含 `$`，
# 不会被 INLINE_FORMULA_PAT 匹配。
_PROTECT_CODE_SPAN = re.compile(r"`[^`\n]+`")
_PROTECT_ESCAPED_DOLLAR = re.compile(r"\\\$")
_PROTECT_KEY_OPEN = "\ue000"
_PROTECT_KEY_CLOSE = "\ue001"
# 行内公式内容守卫（对齐 pandoc 启发式 + C 题经济文本加固）：
_FORMULA_GUARD_CJK = re.compile(r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]")


def _protect_dollar_sensitive(text: str) -> tuple[str, dict[str, str]]:
    """返回（保护后文本, 占位符→还原文本 的映射）。"""
    restore: dict[str, str] = {}

    def _stash(match_text: str, replacement: str) -> str:
        key = f"{_PROTECT_KEY_OPEN}{len(restore)}{_PROTECT_KEY_CLOSE}"
        restore[key] = replacement
        return key

    protected = _PROTECT_CODE_SPAN.sub(lambda m: _stash(m.group(0), m.group(0)), text)
    protected = _PROTECT_ESCAPED_DOLLAR.sub(lambda m: _stash(m.group(0), "$"), protected)
    return protected, restore


def _restore_protected(text: str, restore: dict[str, str]) -> str:
    for key, value in restore.items():
        if key in text:
            text = text.replace(key, value)
    return text


def _is_inline_formula_candidate(content: str, full_text: str, start: int, end: int) -> bool:
    """判定一个 ``$...$`` 匹配是否可信为行内公式（H-10 守卫）：

    - 内容首尾不能是空白（pandoc 规则：`$` 右邻、闭 `$` 左邻须非空格）；
    - 内容含 CJK（汉字/中文标点/全角字符）→ 一定不是 LaTeX 公式；
    - 开 `$` 紧邻前一位是数字、或闭 `$` 紧邻后一位是数字 → 货币金额形态
      （如 "3.5$，4.5$"、"$20,000 和 $30"），拒绝。
    """
    if content[:1].isspace() or content[-1:].isspace():
        return False
    if _FORMULA_GUARD_CJK.search(content):
        return False
    if start > 0 and full_text[start - 1].isdigit():
        return False
    if full_text[end:end + 1].isdigit():
        return False
    return True


def _add_runs_with_inline_formula(paragraph, text: str, font: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    """把含 $...$ 行内公式的文本拆分注入段落：公式走 OMML，其余走普通 run。

    所有需要渲染含公式文本的段落（body/center/heading/list/table cell）都应调用此函数，
    确保 100% 行内公式覆盖。无公式时退化为单个普通 run。

    H-10：切分前先保护 code span 与 ``\\$`` 转义（货币/文件名不再误切进公式），
    再用 _is_inline_formula_candidate 守卫过滤货币形态的伪公式；被守卫拒绝的
    ``$...$`` 保持原样作为正文输出（不吞正文）。
    """
    protected, restore = _protect_dollar_sensitive(text)

    def _plain(segment: str) -> str:
        return clean_inline_markdown(_restore_protected(segment, restore))

    pos = 0
    for m in INLINE_FORMULA_PAT.finditer(protected):
        if not _is_inline_formula_candidate(m.group(1), protected, m.start(), m.end()):
            continue  # 货币/误切形态：跳过，保持正文原样
        if m.start() > pos:
            run = paragraph.add_run(_plain(protected[pos:m.start()]))
            apply_run_font(run, font, size, bold)
        latex = m.group(1).strip()
        if not _append_omml(paragraph, latex):
            run = paragraph.add_run(latex)
            apply_run_font(run, "Cambria Math", size, bold)
        pos = m.end()
    if pos < len(protected):
        run = paragraph.add_run(_plain(protected[pos:]))
        apply_run_font(run, font, size, bold)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_runs_with_inline_formula(paragraph, text, "宋体", 10.5, False)


def add_center_paragraph(document: Document, text: str, font_name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    _add_runs_with_inline_formula(paragraph, text, font_name, size, bold)


def add_block_formula(document: Document, latex: str) -> None:
    """块级公式：居中无缩进段落 + Word 原生 OMML。"""
    latex = latex.strip()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    if not _append_omml(paragraph, latex):
        # 兜底：纯文本（非 OMML），至少不丢内容
        run = paragraph.add_run(latex)
        apply_run_font(run, "Cambria Math", 10.5)


def add_heading(document: Document, text: str, level: int) -> None:
    level = max(1, min(level, 3))
    paragraph = document.add_heading("", level=level)
    paragraph.paragraph_format.first_line_indent = None
    _add_runs_with_inline_formula(paragraph, text, "黑体", {1: 15, 2: 13, 3: 12}[level], True)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(code.rstrip())
    apply_run_font(run, "Consolas", 8.5)


def read_csv_rows(path: Path, max_rows: int = 18, max_cols: int = 8) -> list[list[str]]:
    if not path.exists():
        return []
    encodings = ("utf-8-sig", "utf-8", "gbk")
    for encoding in encodings:
        try:
            with path.open("r", encoding=encoding, newline="") as handle:
                rows = [[str(cell) for cell in row[:max_cols]] for row in csv.reader(handle)]
            return rows[:max_rows]
        except Exception:
            continue
    return []


def add_table_from_rows(document: Document, rows: list[list[str]], caption: str | None = None) -> None:
    if not rows:
        if caption:
            add_center_paragraph(document, caption, bold=True)
        add_body_paragraph(document, "表格数据文件暂不可读取，正式提交前需检查表格索引和源 CSV 文件。")
        return
    if caption:
        add_center_paragraph(document, caption, bold=True)
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            set_cell_margin(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F2F2")
            # 用 cell 默认空 paragraph 注入含公式切分的 runs
            # （不用 cell.text=... ，那会预填 run 并触发 clean_inline_markdown 破坏公式）
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.first_line_indent = None
            cell_para.paragraph_format.space_after = Pt(0)
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            _add_runs_with_inline_formula(cell_para, value, "宋体", 9, row_idx == 0)
    document.add_paragraph()


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    add_table_from_rows(document, rows)


def build_table_lookup(table_index: Any) -> dict[str, dict[str, Any]]:
    lookup = {}
    for item in table_index.get("tables", []) if isinstance(table_index, dict) else []:
        if not isinstance(item, dict):
            continue
        table_id = str(item.get("table_id") or "").strip()
        if table_id:
            lookup[table_id] = item
    return lookup


def build_figure_lookup(figure_index: Any) -> dict[str, dict[str, Any]]:
    lookup = {}
    for item in figure_index.get("figures", []) if isinstance(figure_index, dict) else []:
        if not isinstance(item, dict):
            continue
        figure_id = str(item.get("figure_id") or "").strip()
        if figure_id:
            lookup[figure_id] = item
    return lookup


def add_index_table(document: Document, table_id: str, table_lookup: dict[str, dict[str, Any]]) -> bool:
    """按 table_id 嵌入表格。索引未命中（H-8）时插入可见占位段落而非静默消失。"""
    item = table_lookup.get(table_id)
    if not item:
        add_body_paragraph(document, f"【表格占位：索引未命中 {table_id}，请补齐 table_index.json 条目或数据文件】")
        return False
    rows = read_csv_rows(resolve_path(str(item.get("path") or ""), TABLE_EXT_WHITELIST))
    caption = item.get("caption") or item.get("title") or table_id
    if not str(caption).startswith("表"):
        caption = f"表 {caption}"
    add_table_from_rows(document, rows, caption=str(caption))
    return True


def add_image(document: Document, path: Path, caption: str | None = None) -> bool:
    if not path.exists():
        add_body_paragraph(document, f"图片文件未找到：{rel(path)}。正式提交前需补齐图像文件。")
        return False
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), width=Cm(14.2))
    except Exception:
        try:
            run.add_picture(str(path), width=Cm(12.8))
        except Exception:
            add_body_paragraph(document, f"图片无法插入：{rel(path)}。")
            return False
    if caption:
        add_center_paragraph(document, caption, bold=True)
    return True


def add_index_figure(document: Document, figure_id: str, figure_lookup: dict[str, dict[str, Any]]) -> bool:
    """按 figure_id 嵌入图片。索引未命中（H-8）时插入可见占位段落而非静默消失；
    文件缺失/被安全门禁拒绝时由 add_image 走"图片文件未找到"可见占位。"""
    item = figure_lookup.get(figure_id)
    if not item:
        add_body_paragraph(document, f"【图表占位：索引未命中 {figure_id}，请补齐 figure_index.json 条目或图片文件】")
        return False
    caption = item.get("caption") or item.get("title") or figure_id
    if not str(caption).startswith("图"):
        caption = f"图 {caption}"
    return add_image(document, resolve_path(str(item.get("path") or item.get("expected_path") or ""), IMAGE_EXT_WHITELIST), str(caption))


def source_path() -> Path:
    if SOURCE_FILE.exists():
        return SOURCE_FILE
    return FALLBACK_SOURCE_FILE


def render_markdown(document: Document, text: str, table_lookup: dict[str, dict[str, Any]], figure_lookup: dict[str, dict[str, Any]]) -> dict[str, Any]:
    stats: dict[str, Any] = {"headings": 0, "tables": 0, "figures": 0, "code_blocks": 0}
    # 每次渲染复位全局计数（CR-7 三计数 + H-1/H-8 失败清单）
    _OMML_STATS.update(total=0, ok=0, fallback=0)
    _RENDER_FAILURES.clear()
    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    in_formula = False
    formula_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                stats["code_blocks"] += 1
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        # 多行块级公式：$$ 单独成行 → 收集到下一个 $$ 为止
        if stripped == "$$":
            if in_formula:
                add_block_formula(document, "\n".join(formula_lines))
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
                formula_lines = []
            idx += 1
            continue
        # 单行块级公式 $$...$$（源稿主要用这种）
        single_block = re.fullmatch(r"\$\$(.+)\$\$", stripped)
        if single_block:
            add_block_formula(document, single_block.group(1))
            idx += 1
            continue
        if in_formula:
            formula_lines.append(stripped)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        table_marker = re.fullmatch(r"\[\[TABLE:([A-Za-z0-9_\-]+)\]\]", stripped, flags=re.IGNORECASE)
        if table_marker:
            if add_index_table(document, table_marker.group(1), table_lookup):
                stats["tables"] += 1
            else:  # H-8：索引未命中/数据不可读 → 已插可见占位，这里补记 failure
                _RENDER_FAILURES.append(f"[[TABLE:{table_marker.group(1)}]] 未成功嵌入（索引未命中或数据文件不可读）")
            idx += 1
            continue

        figure_marker = re.fullmatch(r"\[\[FIGURE:([A-Za-z0-9_\-]+)\]\]", stripped, flags=re.IGNORECASE)
        if figure_marker:
            if add_index_figure(document, figure_marker.group(1), figure_lookup):
                stats["figures"] += 1
            else:  # H-8：同上，不再静默消失
                _RENDER_FAILURES.append(f"[[FIGURE:{figure_marker.group(1)}]] 未成功嵌入（索引未命中或图片文件缺失）")
            idx += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            caption = image_match.group(1).strip()
            path = resolve_path(image_match.group(2).strip(), IMAGE_EXT_WHITELIST)
            if add_image(document, path, caption or None):
                stats["figures"] += 1
            idx += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = [stripped]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|") and "|" in lines[idx].strip()[1:]:
                table_lines.append(lines[idx].strip())
                idx += 1
            add_markdown_table(document, table_lines)
            stats["tables"] += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            add_heading(document, heading.group(2), level)
            stats["headings"] += 1
            idx += 1
            continue

        numbered_heading = re.match(r"^((?:\d+\.){0,2}\d+)\s+(.+)$", stripped)
        if numbered_heading and len(stripped) <= 80:
            level = numbered_heading.group(1).count(".") + 1
            add_heading(document, stripped, level)
            stats["headings"] += 1
            idx += 1
            continue

        list_item = re.match(r"^[-*]\s+(.+)$", stripped)
        if list_item:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.25)
            bullet_run = paragraph.add_run("• ")
            apply_run_font(bullet_run, "宋体", 10.5)
            _add_runs_with_inline_formula(paragraph, list_item.group(1), "宋体", 10.5, False)
            idx += 1
            continue

        add_body_paragraph(document, stripped)
        idx += 1

    if code_lines:
        add_code_block(document, "\n".join(code_lines))
        stats["code_blocks"] += 1
    if formula_lines:
        add_block_formula(document, "\n".join(formula_lines))
    # CR-7：合并 OMML 三计数与渲染失败清单进报告 stats
    stats["omml_total"] = _OMML_STATS["total"]
    stats["omml_ok"] = _OMML_STATS["ok"]
    stats["omml_fallback"] = _OMML_STATS["fallback"]
    stats["failures"] = list(_RENDER_FAILURES)
    return stats


def write_report(stats: dict[str, Any], source: Path, outline: Any) -> str:
    """写格式化报告，返回状态字符串（GENERATED / DEGRADED）。"""
    omml_total = int(stats.get("omml_total", 0))
    omml_ok = int(stats.get("omml_ok", 0))
    omml_fallback = int(stats.get("omml_fallback", 0))
    failures = list(stats.get("failures", []))
    # CR-7：公式链退化（依赖缺失或有回退）时状态降级，不再零上报
    degraded = (omml_fallback > 0) or (not _FORMULA_AVAILABLE)
    status = "DEGRADED" if degraded else "GENERATED"
    lines = [
        "# Formal DOCX Formatting Report",
        "",
        f"- Status: `{status}`",
        f"- Generated at: `{datetime.now().isoformat(timespec='seconds')}`",
        f"- Source: `{rel(source)}`",
        f"- Output: `{rel(DOCX_FILE)}`",
        f"- Outline: `{rel(OUTLINE_FILE) if OUTLINE_FILE.exists() else 'missing'}`",
        f"- Title: `{outline.get('title', '') if isinstance(outline, dict) else ''}`",
        f"- Headings: `{stats.get('headings', 0)}`",
        f"- Tables inserted: `{stats.get('tables', 0)}`",
        f"- Figures inserted: `{stats.get('figures', 0)}`",
        f"- Code blocks: `{stats.get('code_blocks', 0)}`",
        f"- OMML formulas (total): `{omml_total}`",
        f"- OMML formulas (ok, Word 原生公式): `{omml_ok}`",
        f"- OMML formulas (fallback, Cambria Math 纯文本): `{omml_fallback}`",
        f"- Formula engine: `{'latex_to_omml via pandoc' if _FORMULA_AVAILABLE else 'UNAVAILABLE'}`",
        "- Render QA: `render_skipped`",
        "",
    ]
    if not _FORMULA_AVAILABLE:
        lines.append(
            "> ⚠ **DEGRADED 退化原因**：无法加载 `docx-editor-cn/scripts/formula.py::latex_to_omml`"
            "（通常为 pandoc 缺失或依赖不可达），**全部公式退化为 Cambria Math 纯文本**——"
            "Word 中双击不可进公式编辑器。请安装 pandoc ≥ 2.0 后重跑本脚本，"
            "并确认 `_DOCX_EDITOR_SCRIPTS` 定位成功。"
        )
        lines.append("")
    elif omml_fallback > 0:
        lines.append(
            f"> ⚠ **DEGRADED 退化原因**：{omml_fallback}/{omml_total} 个公式 OMML 转换或注入失败，"
            "退化为 Cambria Math 纯文本（失败明细见运行 stderr 的 `[formula]` 日志）。"
            "正式交付前应使 fallback 归零。"
        )
        lines.append("")
    if failures:
        lines.extend(["## Failures", ""])
        lines.extend(f"- {item}" for item in failures)
        lines.append("")
    lines.append(
        "LibreOffice 渲染不是本脚本的强依赖；若本机 LibreOffice 可用，可在最终交付前另行渲染 PNG/PDF 做视觉检查。"
    )
    REPORT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return status


def check_evidence_gate() -> tuple[bool, str]:
    """Return (passed, reason). passed=True only when report exists and status==PASS."""
    if not EVIDENCE_GATE_REPORT.exists():
        return False, f"未找到证据门禁报告：{rel(EVIDENCE_GATE_REPORT)}。请先运行 quality-assurance-auditor/scripts/evidence_gate.py。"
    try:
        data = json.loads(EVIDENCE_GATE_REPORT.read_text(encoding="utf-8"))
    except Exception as exc:
        return False, f"证据门禁报告无法解析：{type(exc).__name__}: {exc}"
    status = str(data.get("status") or "").strip().upper()
    if status != "PASS":
        return False, f"证据门禁状态为 `{status or 'UNKNOWN'}`，正式 Word 不得生成。请先补齐证据并重跑 evidence_gate.py。"
    return True, ""


def main() -> int:
    configure_utf8_stdio()
    parser = argparse.ArgumentParser(description="Format the formal paper DOCX from final_paper_source.md.")
    parser.add_argument(
        "--allow-draft",
        action="store_true",
        help="证据门禁未通过时仍生成草稿 Word（写入 final_paper_draft.docx + format_draft_report.md），不会覆盖正式产物。",
    )
    args = parser.parse_args()

    global DOCX_FILE, REPORT_MD

    gate_passed, gate_reason = check_evidence_gate()
    draft_mode = False
    if not gate_passed:
        if not args.allow_draft:
            print("[FORMAT BLOCKED] 证据门禁未通过，禁止生成正式 final_paper.docx。", file=sys.stderr)
            print(f"  原因：{gate_reason}", file=sys.stderr)
            print("  如需先看排版草稿，请加 --allow-draft，会写入 final_paper_draft.docx，不会污染正式产物。", file=sys.stderr)
            return 2
        draft_mode = True
        DOCX_FILE = DOCX_FILE_DRAFT
        REPORT_MD = REPORT_MD_DRAFT
        print(f"[DRAFT MODE] 证据门禁未通过：{gate_reason}")
        print(f"[DRAFT MODE] 将写入草稿 Word：{rel(DOCX_FILE)}（不会覆盖正式 final_paper.docx）")

    source = source_path()
    if not source.exists():
        print(f"缺少正式论文 Markdown：{rel(SOURCE_FILE)}", file=sys.stderr)
        return 1

    outline = load_json(OUTLINE_FILE)
    table_index = load_json(TABLE_INDEX_FILE)
    figure_index = load_json(FIGURE_INDEX_FILE)

    document = Document()
    configure_document(document)

    text = source.read_text(encoding="utf-8")
    stats = render_markdown(document, text, build_table_lookup(table_index), build_figure_lookup(figure_index))
    DOCX_FILE.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_FILE)
    status = write_report(stats, source, outline)
    label = "草稿 Word" if draft_mode else "正式 Word"
    print(f"{label}已生成：{rel(DOCX_FILE)}")
    print(f"格式化报告已生成：{rel(REPORT_MD)}（Status: {status}）")
    if status == "DEGRADED":
        # CR-7：退化必须在控制台同样显著可见，不再静默 GENERATED
        if not _FORMULA_AVAILABLE:
            print("[DEGRADED] 公式引擎不可用（latex_to_omml/pandoc 缺失），全部公式退化为 Cambria Math 纯文本，非 Word 原生 OMML。")
        else:
            print(f"[DEGRADED] {_OMML_STATS['fallback']}/{_OMML_STATS['total']} 个公式 OMML 转换失败，退化为纯文本。")
    if stats.get("failures"):
        print(f"[RENDER FAILURES] {len(stats['failures'])} 条（路径拒绝/索引未命中/文件缺失），详见报告 Failures 节。")
    if draft_mode:
        print("[DRAFT MODE] 该文件不是最终稿；正式提交前必须先通过证据门禁，再不带 --allow-draft 重跑本脚本。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
