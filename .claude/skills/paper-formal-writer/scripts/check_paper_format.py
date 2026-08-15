import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from _project_root import ProjectRootNotFoundError, find_project_root

# M-13（未收口 #9）锚定统一：由文件位置上溯定位项目根，替代 Path.cwd()——
# 从任意 cwd 运行都锚定真实 paper_output/；错误 cwd 不再读错源稿/写错报告/
# 凭空创建 paper_output 树；脚本被复制到项目外时启动即报错（见 _project_root.py）。
try:
    BASE_DIR = find_project_root()
except ProjectRootNotFoundError as exc:
    print(f"[ANCHOR FAILED] {exc}", file=sys.stderr)
    raise SystemExit(2)
OUTPUT_DIR = BASE_DIR / "paper_output"
SOURCE_FILE = OUTPUT_DIR / "final_paper_source.md"
FALLBACK_SOURCE_FILE = OUTPUT_DIR / "final_paper.md"
DOCX_FILE = OUTPUT_DIR / "final_paper.docx"
OUTLINE_FILE = OUTPUT_DIR / "plan" / "paper_outline.json"
FIGURE_INDEX_FILE = OUTPUT_DIR / "figure_index.json"
FIGURE_INDEX_FALLBACK = OUTPUT_DIR / "figures" / "figure_index.json"
TABLE_INDEX_FILE = OUTPUT_DIR / "tables" / "table_index.json"
TABLE_INDEX_FALLBACK = OUTPUT_DIR / "table_index.json"
# format_formal_docx.py 的渲染报告（Status: GENERATED/DEGRADED + OMML fallback 计数）
RENDER_REPORT_FILE = OUTPUT_DIR / "format_check_report.md"
# 双写冲突修复（MEDIUM）：format_check_report.md 由 format_formal_docx.py 独占，
# 本门禁输出改独立文件名 format_check_gate_report.md。
# format_check_report.json 保持原名——workflow_guard.py check_s8 按该名消费 status。
REPORT_MD = OUTPUT_DIR / "format_check_gate_report.md"
REPORT_JSON = OUTPUT_DIR / "format_check_report.json"

PLACEHOLDERS = [
    "内容生成中",
    "关键词1",
    "论文题目缺失",
    "TODO",
    "待补",
    "{{",
    "}}",
]

REQUIRED_SECTIONS = [
    "摘要",
    "关键词",
    "1 问题重述",
    "2 问题分析",
    "3 模型假设",
    "4 符号说明",
    "5 模型建立与求解",
    "6 灵敏度分析",
    "7 模型检验",
    "8 模型评价",
    "9 结论",
    "参考文献",
]

# 中文数字映射，用于匹配 docx 中的章节标题
CN_NUMBERS = {
    '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
    '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
}


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def load_json(path: Path) -> Any:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"__error__": str(exc)}


def rel(path: Path) -> str:
    try:
        return path.relative_to(BASE_DIR).as_posix()
    except ValueError:
        return path.as_posix()


def source_path() -> Path:
    if SOURCE_FILE.exists():
        return SOURCE_FILE
    return FALLBACK_SOURCE_FILE


def compact_text(text: str) -> str:
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[[^\]]+\]\([^)]+\)", "", text)
    text = re.sub(r"`{1,3}", "", text)
    text = re.sub(r"[#>*_\-|$`{}\[\]():;,.，。；：！？、\s]", "", text)
    return text


def char_count(text: str) -> dict[str, int]:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    nonspace = len(re.sub(r"\s+", "", text))
    content = len(compact_text(text))
    return {"cjk": cjk, "nonspace": nonspace, "content": content}


def has_required_section(text: str, label: str) -> bool:
    if label in ("摘要", "关键词", "附录", "参考文献"):
        return re.search(rf"(^|\n)\s*(?:#+\s*)?(?:\*\*)?{re.escape(label)}", text) is not None
    if " " not in label:
        return re.search(rf"(^|\n)\s*(?:#+\s*)?{re.escape(label)}", text) is not None
    number, title = label.split(" ", 1)
    # 提取标题关键词（去掉"的"等虚词）
    title_keywords = [w for w in title if w not in ("的", "与", "和", "及")]
    # 匹配阿拉伯数字：1 问题重述 / 1. 问题重述 / 1、问题重述
    if re.search(rf"(^|\n)\s*#*\s*{re.escape(number)}[.、\s]+", text):
        # 检查同一行是否包含标题关键词
        for line in text.split("\n"):
            if re.search(rf"^\s*#*\s*{re.escape(number)}[.、\s]+", line):
                if any(kw in line for kw in title_keywords):
                    return True
    # 匹配中文数字：一、问题重述 / 一 问题重述
    for cn, ar in CN_NUMBERS.items():
        if ar == number:
            if re.search(rf"(^|\n)\s*#*\s*{re.escape(cn)}[、\s]+", text):
                for line in text.split("\n"):
                    if re.search(rf"^\s*#*\s*{re.escape(cn)}[、\s]+", line):
                        if any(kw in line for kw in title_keywords):
                            return True
    return False


def natural_q_key(qid: str) -> tuple[int, str]:
    match = re.search(r"\d+", qid)
    if match:
        return (int(match.group()), qid)
    return (10_000, qid)


def qids_from_outline(outline: Any) -> list[str]:
    if isinstance(outline, dict) and isinstance(outline.get("questions"), list):
        qids = [str(item.get("question_id") or "").strip() for item in outline["questions"] if isinstance(item, dict)]
        qids = [qid for qid in qids if qid]
        if qids:
            return sorted(set(qids), key=natural_q_key)
    model_route = load_json(OUTPUT_DIR / "plan" / "model_route.json")
    qids = []
    for item in model_route.get("questions", []) if isinstance(model_route, dict) else []:
        if isinstance(item, dict) and item.get("question_id"):
            qids.append(str(item["question_id"]))
    return sorted(set(qids), key=natural_q_key)


def load_first(*paths: Path) -> Any:
    """按候选路径加载 JSON（索引文件根/子目录双布局兼容）。

    M-13-LOW（load_first 无 mtime 裁决）修复：旧逻辑"第一个存在且合法者胜出"，
    双布局两份索引并存时固定取先传的根路径——根目录残留的陈旧索引会压过
    子目录新索引。现改为：全部存在且可解析的候选中取 mtime 最新者；
    仅一份时行为与旧版一致（存在的第一份即胜出）。
    """
    candidates: list[tuple[float, dict[str, Any]]] = []
    for path in paths:
        if not path.exists():
            continue
        data = load_json(path)
        if isinstance(data, dict) and "__error__" not in data:
            try:
                mtime = path.stat().st_mtime
            except OSError:
                mtime = -1.0  # stat 竞态失败时退化为最低优先，不中断加载
            candidates.append((mtime, data))
    if not candidates:
        return {}
    # max 在 mtime 并列时保留先出现的候选 → 与旧版"先根后子"顺序一致（稳定）
    return max(candidates, key=lambda pair: pair[0])[1]


def item_id(item: dict[str, Any], *id_keys: str) -> str:
    for key in id_keys:
        value = item.get(key)
        if value:
            return str(value)
    return "?"


def index_items(data: Any, key: str, *id_keys: str) -> list[dict[str, Any]]:
    """取索引条目。id 键双读（对齐 evidence_gate）：figure 条目用 figure_id，
    table 条目实际可能是 table_id 或 id——单键读取会把整个维度架空。"""
    if not isinstance(data, dict):
        return []
    items = data.get(key, [])
    if not isinstance(items, list):
        return []
    return [item for item in items if isinstance(item, dict) and any(item.get(k) for k in id_keys)]


def referenced(text: str, item: dict[str, Any], kind: str, *id_keys: str) -> bool:
    """正文是否引用了该图/表。

    除 id/标题/文件名 stem 直接命中外，支持中文编号引用「图 N」「表 N」
    （N 取 id 字段中的数字，如 F1/T1 → 1）——正文交叉引用的实际形态。
    """
    candidates = [str(item.get(key) or "") for key in id_keys]
    candidates.append(str(item.get("title") or ""))
    stem_source = item.get("path") or item.get("expected_path") or ""
    if stem_source:
        candidates.append(Path(str(stem_source)).stem)
    candidates = [candidate.strip() for candidate in candidates if candidate and candidate.strip()]
    if any(candidate in text for candidate in candidates):
        return True

    label = "图" if kind == "figure" else "表"
    for key in (*id_keys, "number"):
        match = re.search(r"(\d+)", str(item.get(key) or ""))
        if match and re.search(rf"{label}\s*{match.group(1)}(?!\d)", text):
            return True
    return False


def check_docx_structure(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"exists": False}
    try:
        doc = Document(str(path))
        headings = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.style and paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ]
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        return {
            "exists": True,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "image_count": image_count,
            "heading_count": len(headings),
            "sample_headings": headings[:12],
        }
    except Exception as exc:
        return {"exists": True, "error": str(exc)}


def check_docx_fonts(path: Path) -> dict[str, Any]:
    """检查 Word 文档字体是否符合国赛论文规范。

    规范：
    - 正文：宋体（中文）+ Times New Roman（英文），小四号（12pt）
    - 一级标题：黑体（中文），三号（16pt），加粗
    - 二级标题：黑体（中文），四号（14pt），加粗
    - 三级标题：黑体（中文），小四号（12pt），加粗
    - 表格：宋体（中文），五号（10.5pt）
    """
    if not path.exists():
        return {"exists": False, "failures": ["Word 文件不存在"], "warnings": []}

    try:
        doc = Document(str(path))
    except Exception as exc:
        return {"exists": True, "failures": [f"Word 文件无法读取：{exc}"], "warnings": []}

    failures = []
    warnings = []
    EA_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"

    # 检查正文段落字体
    body_font_issues = []
    body_size_issues = []
    body_count = 0

    # 收集所有表格内的段落，用于排除
    table_paragraphs = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_paragraphs.add(id(para))

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        # 跳过标题段落
        if para.style and para.style.name.startswith("Heading"):
            continue
        # 跳过表格内的段落
        if id(para) in table_paragraphs:
            continue
        body_count += 1
        for run in para.runs:
            if not run.text.strip():
                continue
            # 检查字号
            if run.font.size is not None:
                size_pt = run.font.size.pt
                if size_pt < 11 or size_pt > 13:  # 小四=12pt，允许±1pt误差
                    body_size_issues.append(f"正文段落字号异常：{size_pt:.1f}pt（应为12pt）")
            # 检查中文字体
            rPr = run.element.rPr
            if rPr is not None:
                rFonts = rPr.rFonts
                if rFonts is not None:
                    ea_font = rFonts.get(EA_NS)
                    if ea_font and ea_font not in ("宋体", "SimSun"):
                        body_font_issues.append(f"正文中文字体不是宋体：{ea_font}")
            break  # 只检查每个段落的第一个 run
        if len(body_font_issues) >= 3:
            break  # 最多报3个

    if body_font_issues:
        warnings.extend(body_font_issues)
    if body_size_issues:
        warnings.extend(body_size_issues[:3])

    # 检查标题字体
    heading_font_issues = []
    heading_size_issues = []
    for para in doc.paragraphs:
        if not para.style or not para.style.name.startswith("Heading"):
            continue
        level = int(para.style.name.replace("Heading ", "1"))
        for run in para.runs:
            if not run.text.strip():
                continue
            # 检查字号
            expected_sizes = {1: 16, 2: 14, 3: 12}
            if level in expected_sizes and run.font.size is not None:
                size_pt = run.font.size.pt
                expected = expected_sizes[level]
                if abs(size_pt - expected) > 1.5:
                    heading_size_issues.append(
                        f"Heading {level} 字号异常：{size_pt:.1f}pt（应为{expected}pt）"
                    )
            # 检查中文字体
            rPr = run.element.rPr
            if rPr is not None:
                rFonts = rPr.rFonts
                if rFonts is not None:
                    ea_font = rFonts.get(EA_NS)
                    if ea_font and ea_font not in ("黑体", "SimHei"):
                        heading_font_issues.append(
                            f"Heading {level} 中文字体不是黑体：{ea_font}"
                        )
            break
        if len(heading_font_issues) >= 3:
            break

    if heading_font_issues:
        failures.extend(heading_font_issues)
    if heading_size_issues:
        failures.extend(heading_size_issues[:3])

    # 检查默认样式字体
    try:
        normal_style = doc.styles["Normal"]
        if normal_style.font.name and normal_style.font.name not in (
            "Times New Roman", "宋体", "SimSun"
        ):
            warnings.append(f"默认正文字体不是 Times New Roman：{normal_style.font.name}")
    except Exception:
        pass

    return {
        "exists": True,
        "body_paragraphs_checked": body_count,
        "heading_font_issues": heading_font_issues,
        "heading_size_issues": heading_size_issues,
        "body_font_issues": body_font_issues,
        "body_size_issues": body_size_issues,
        "failures": failures,
        "warnings": warnings,
    }


def markdown_heading_count(text: str) -> int:
    return len(re.findall(r"(^|\n)\s*#{1,6}\s+\S+", text))


def visual_qa_failures(
    docx_structure: dict[str, Any],
    source_heading_count: int,
    figure_count: int,
    table_count: int,
) -> tuple[list[str], list[str]]:
    failures: list[str] = []
    warnings: list[str] = []
    if not docx_structure.get("exists") or docx_structure.get("error"):
        return failures, warnings

    paragraph_count = int(docx_structure.get("paragraph_count") or 0)
    heading_count = int(docx_structure.get("heading_count") or 0)
    docx_table_count = int(docx_structure.get("table_count") or 0)
    image_count = int(docx_structure.get("image_count") or 0)

    if paragraph_count < 10:
        failures.append(f"Word 段落数量异常偏少：{paragraph_count} < 10")
    if source_heading_count > 0 and heading_count == 0:
        failures.append("Word 中没有可识别标题样式，标题结构可能未正确写入。")
    elif source_heading_count > 0 and heading_count < max(1, source_heading_count // 2):
        warnings.append(f"Word 标题数量明显少于 Markdown 标题：{heading_count} < {source_heading_count}")

    if figure_count > 0 and image_count == 0:
        failures.append("figure_index.json 有图片计划，但 Word 中没有图片。")
    elif image_count < figure_count:
        # H-9：缺嵌从 warning 升级为 failure——"N 图全嵌"必须是机制保证而非运气
        failures.append(f"Word 图片数量少于 figure_index.json：{image_count} < {figure_count}")

    if table_count > 0 and docx_table_count == 0:
        failures.append("table_index.json 有表格计划，但 Word 中没有表格。")
    elif docx_table_count < table_count:
        warnings.append(f"Word 表格数量少于 table_index.json：{docx_table_count} < {table_count}")

    return failures, warnings


def extract_docx_text(path: Path) -> str:
    """从 docx 文件提取全部文本（段落+表格）。"""
    if not path.exists():
        return ""
    try:
        doc = Document(str(path))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            texts.append(para.text)
        return "\n".join(texts)
    except Exception:
        return ""


def check_render_report(path: Path) -> dict[str, Any]:
    """读 format_formal_docx 的渲染报告（CR-7 证据接进门禁）。

    - Status: DEGRADED → failure（公式链退化：引擎不可用或有回退）
    - OMML fallback > 0 → failure（有公式退化为 Cambria Math 纯文本）
    - 报告缺失 → warning（公式链退化不可见，不能装作已核验）
    """
    result: dict[str, Any] = {"exists": False, "status": "MISSING", "omml_fallback": None,
                              "failures": [], "warnings": []}
    if not path.exists():
        result["warnings"].append(
            f"未找到 format_formal_docx 渲染报告（{rel(path)}），公式链退化不可见——"
            "请先运行 format_formal_docx.py 再过本门禁"
        )
        return result
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        result["warnings"].append(f"渲染报告无法读取：{exc}")
        return result

    status_match = re.search(r"-\s*Status:\s*`?([A-Za-z]+)`?", content)
    result["exists"] = True
    result["status"] = status_match.group(1) if status_match else "UNKNOWN"
    fallback_match = re.search(r"OMML formulas \(fallback[^)]*\):\s*`?(\d+)", content)
    if fallback_match:
        result["omml_fallback"] = int(fallback_match.group(1))

    if result["status"] == "DEGRADED":
        result["failures"].append(
            f"format_formal_docx 渲染报告状态 DEGRADED：公式链退化（详见 {rel(path)}），正式交付前必须修复"
        )
    if result["omml_fallback"] is not None and result["omml_fallback"] > 0:
        result["failures"].append(
            f"OMML 公式回退 {result['omml_fallback']} 个（Cambria Math 纯文本，非 Word 原生公式），正式交付前应归零"
        )
    if result["status"] not in ("GENERATED", "DEGRADED"):
        result["warnings"].append(f"渲染报告状态异常：{result['status']}")
    return result


def evaluate(source_override: Path | None = None) -> dict[str, Any]:
    source = source_override if source_override is not None else source_path()
    outline = load_json(OUTLINE_FILE)
    figure_index = load_first(FIGURE_INDEX_FILE, FIGURE_INDEX_FALLBACK)
    table_index = load_first(TABLE_INDEX_FILE, TABLE_INDEX_FALLBACK)
    failures: list[str] = []
    warnings: list[str] = []

    # 优先从 docx 提取文本，fallback 到 markdown；
    # 显式指定 --source 时改用该文件作为正文文本来源
    docx_text = extract_docx_text(DOCX_FILE)
    if source_override is not None:
        if source.exists():
            try:
                text = source.read_text(encoding="utf-8")
                source_label = "markdown"
            except Exception as exc:
                failures.append(f"指定的论文源文件无法读取：{rel(source)}（{exc}）")
                text = ""
                source_label = "none"
        else:
            failures.append(f"指定的论文源文件不存在：{rel(source)}")
            text = ""
            source_label = "none"
    elif docx_text:
        text = docx_text
        source_label = "docx"
    elif source.exists():
        text = source.read_text(encoding="utf-8")
        source_label = "markdown"
    else:
        failures.append(f"缺少正式论文文件：{rel(DOCX_FILE)} 和 {rel(SOURCE_FILE)}")
        text = ""
        source_label = "none"

    # 实伤-2 防线：全文 U+FFFD 扫描——替换符意味着编码链路已丢字，直接判 FAIL
    fffd_count = text.count("\ufffd")
    if fffd_count:
        failures.append(f"正文含 U+FFFD 替换符（编码丢字信号）：{fffd_count} 处，需人工核对源稿补全")

    counts = char_count(text)
    # target_words 兼容两种形态（C-①）：
    #   int（真实 outline 即 int 18000，单值即下限） / dict（{"min":..., "max":...}）
    raw_target = outline.get("target_words") if isinstance(outline, dict) else None
    if isinstance(raw_target, dict):
        min_words = int(raw_target.get("min", 18000) or 18000)
        max_words = int(raw_target.get("max", 25000) or 25000)
    elif isinstance(raw_target, int) and not isinstance(raw_target, bool) and raw_target > 0:
        min_words = raw_target
        max_words = max(25000, raw_target)
    else:
        min_words, max_words = 18000, 25000
    # 使用 nonspace 计数（包含标点），更适合中文学术论文
    # 对于 docx 来源，使用更宽松的计数（包含空格）
    if source_label == "docx":
        effective_count = len(text)  # docx 总字符数（含空格）
    else:
        effective_count = counts["nonspace"]
    if effective_count < min_words:
        failures.append(f"正文有效字数不足：{effective_count} < {min_words}")
    if effective_count > max_words:
        warnings.append(f"正文有效字数超过建议上限：{effective_count} > {max_words}")

    missing_sections = [label for label in REQUIRED_SECTIONS if not has_required_section(text, label)]
    for label in missing_sections:
        failures.append(f"缺少正式论文结构：{label}")

    # 检查 5.1 / 5.1.1 / 5.1.2（支持中文编号：5.1 或 五、1 或 5.1）
    if not re.search(r"(^|\n)\s*#*\s*(5\.1|五[、.\s]*1)\s+", text):
        failures.append("缺少 5.1 问题一模型章节。")
    if not re.search(r"(^|\n)\s*#*\s*(5\.1\.1|5\.1\.1)\s+", text):
        # docx 中可能是「5.1.1 模型建立」或「5.1.1 模型建立」
        if not re.search(r"(^|\n)\s*#*\s*5\.1\.1\s+", text):
            failures.append("缺少 5.1.1 三级标题。")
    if not re.search(r"(^|\n)\s*#*\s*5\.1\.2\s+", text):
        failures.append("缺少 5.1.2 三级标题。")

    question_reports = []
    for index, qid in enumerate(qids_from_outline(outline), start=1):
        q_failures: list[str] = []
        section = f"5.{index}"
        section_pattern = rf"(^|\n)\s*#*\s*{re.escape(section)}\s+"
        if not re.search(section_pattern, text):
            q_failures.append(f"缺少 {section} 对应 {qid} 的模型章节")
        # 检查子章节（灵活匹配，不要求所有子章节都存在）
        subsections_found = 0
        for suffix in ("1", "2", "3", "4", "5"):
            if re.search(rf"(^|\n)\s*#*\s*{re.escape(section + '.' + suffix)}\s+", text):
                subsections_found += 1
        if subsections_found == 0:
            q_failures.append(f"缺少 {section} 的任何子章节")
        if not re.search(rf"{qid}|问题[一二三四五六七八九十{index}]", text):
            q_failures.append(f"正文未明确回扣 {qid}")
        question_reports.append({"question_id": qid, "status": "FAIL" if q_failures else "PASS", "failures": q_failures})
        failures.extend(q_failures)

    figures = index_items(figure_index, "figures", "figure_id", "id")
    tables = index_items(table_index, "tables", "table_id", "id")
    missing_figures = [item_id(item, "figure_id", "id") for item in figures
                       if not referenced(text, item, "figure", "figure_id", "id")]
    missing_tables = [item_id(item, "table_id", "id") for item in tables
                      if not referenced(text, item, "table", "table_id", "id")]
    for figure_id in missing_figures:
        failures.append(f"figure_index.json 中的图片未在正文引用：{figure_id}")
    for table_id in missing_tables:
        failures.append(f"table_index.json 中的表格未在正文引用：{table_id}")
    if len(figures) < 5:
        warnings.append(f"图数量少于展示样例建议值：{len(figures)} < 5")
    if len(tables) < 5:
        warnings.append(f"表数量少于展示样例建议值：{len(tables)} < 5")

    # 检查占位符，但排除LaTeX公式中的内容
    # 移除LaTeX公式块后再检查占位符
    # 使用更宽泛的匹配来移除所有 LaTeX 公式内容
    text_without_latex = re.sub(r'\${1,2}[^$]+\${1,2}', '', text, flags=re.DOTALL)
    for placeholder in PLACEHOLDERS:
        if placeholder in text_without_latex:
            failures.append(f"存在占位符或待补文本：{placeholder}")

    if "参考文献" in text and len(re.findall(r"\[\d+\]", text)) < 3:
        warnings.append("参考文献条目少于 3 条，建议补充权威来源。")

    docx_structure = check_docx_structure(DOCX_FILE)
    if not docx_structure.get("exists"):
        failures.append(f"缺少正式 Word 文件：{rel(DOCX_FILE)}")
    elif docx_structure.get("error"):
        failures.append(f"Word 文件无法读取：{docx_structure['error']}")

    source_heading_count = markdown_heading_count(text)
    visual_failures, visual_warnings = visual_qa_failures(docx_structure, source_heading_count, len(figures), len(tables))
    failures.extend(visual_failures)
    warnings.extend(visual_warnings)

    # 字体检查
    font_check = check_docx_fonts(DOCX_FILE)
    failures.extend(font_check.get("failures", []))
    warnings.extend(font_check.get("warnings", []))

    # 渲染报告检查（format_formal_docx 的 CR-7 证据：DEGRADED / OMML fallback）
    render_check = check_render_report(RENDER_REPORT_FILE)
    failures.extend(render_check.get("failures", []))
    warnings.extend(render_check.get("warnings", []))

    return {
        "schema_version": "1.0",
        "generated_by": "paper-formal-writer/scripts/check_paper_format.py",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "PASS" if not failures else "FAIL",
        "source": source_label,
        "source_file": rel(DOCX_FILE) if source_label == "docx" else rel(source),
        "docx": rel(DOCX_FILE),
        "counts": counts,
        "target_words": {"min": min_words, "max": max_words},
        "render_check": {
            "report": rel(RENDER_REPORT_FILE),
            "exists": render_check.get("exists", False),
            "status": render_check.get("status"),
            "omml_fallback": render_check.get("omml_fallback"),
        },
        "question_reports": question_reports,
        "figure_count": len(figures),
        "table_count": len(tables),
        "missing_figures": missing_figures,
        "missing_tables": missing_tables,
        "source_heading_count": source_heading_count,
        "docx_structure": docx_structure,
        "visual_qa": {
            "status": "PASS" if not visual_failures else "FAIL",
            "failures": visual_failures,
            "warnings": visual_warnings,
        },
        "font_check": {
            "status": "PASS" if not font_check.get("failures") else "FAIL",
            "heading_font_issues": font_check.get("heading_font_issues", []),
            "heading_size_issues": font_check.get("heading_size_issues", []),
            "body_font_issues": font_check.get("body_font_issues", []),
            "body_size_issues": font_check.get("body_size_issues", []),
        },
        "failures": failures,
        "warnings": warnings,
    }


def write_reports(report: dict[str, Any]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Formal Paper Format Check Report",
        "",
        f"- Status: `{report['status']}`",
        f"- Generated at: `{report['generated_at']}`",
        f"- Source: `{report['source']}` ({report.get('source_file', 'N/A')})",
        f"- DOCX: `{report['docx']}`",
        f"- Effective chars: `{report['counts']['content']}`",
        f"- CJK chars: `{report['counts']['cjk']}`",
        f"- Figures in index: `{report['figure_count']}`",
        f"- Tables in index: `{report['table_count']}`",
        "",
    ]
    if report["failures"]:
        lines.append("## Failures")
        lines.extend(f"- {item}" for item in report["failures"])
        lines.append("")
    if report["warnings"]:
        lines.append("## Warnings")
        lines.extend(f"- {item}" for item in report["warnings"])
        lines.append("")
    lines.append("## Question Coverage")
    for item in report["question_reports"]:
        lines.append(f"- {item['question_id']}: `{item['status']}`")
        for failure in item["failures"]:
            lines.append(f"  - {failure}")
    lines.append("")
    lines.append("## DOCX Structure")
    for key, value in report["docx_structure"].items():
        lines.append(f"- {key}: `{value}`")
    lines.append("")
    lines.append("## Visual QA")
    lines.append(f"- status: `{report['visual_qa']['status']}`")
    lines.append(f"- source_heading_count: `{report['source_heading_count']}`")
    for failure in report["visual_qa"]["failures"]:
        lines.append(f"- failure: {failure}")
    for warning in report["visual_qa"]["warnings"]:
        lines.append(f"- warning: {warning}")
    lines.append("")
    lines.append("## Font Check")
    lines.append(f"- status: `{report['font_check']['status']}`")
    lines.append(f"- heading_font_issues: `{len(report['font_check']['heading_font_issues'])}`")
    lines.append(f"- heading_size_issues: `{len(report['font_check']['heading_size_issues'])}`")
    lines.append(f"- body_font_issues: `{len(report['font_check']['body_font_issues'])}`")
    lines.append(f"- body_size_issues: `{len(report['font_check']['body_size_issues'])}`")
    for issue in report["font_check"]["heading_font_issues"] + report["font_check"]["heading_size_issues"]:
        lines.append(f"- failure: {issue}")
    for issue in report["font_check"]["body_font_issues"] + report["font_check"]["body_size_issues"]:
        lines.append(f"- warning: {issue}")
    lines.append("")
    lines.append("## Render Check (format_formal_docx)")
    render = report.get("render_check", {})
    lines.append(f"- report: `{render.get('report', 'N/A')}`")
    lines.append(f"- exists: `{render.get('exists', False)}`")
    lines.append(f"- status: `{render.get('status', 'N/A')}`")
    lines.append(f"- omml_fallback: `{render.get('omml_fallback', 'N/A')}`")
    REPORT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="check_paper_format.py",
        description=(
            "正式论文格式门禁：字数、章节结构、图表引用、占位符、"
            "Word 结构与字体规范检查。"
        ),
        epilog=(
            "零参数运行保持原行为：优先从 paper_output/final_paper.docx 提取正文，"
            "否则读取 paper_output/final_paper_source.md（缺省回退 final_paper.md）。"
            "报告写入 paper_output/format_check_gate_report.md（md 与 format_formal_docx 的"
            " format_check_report.md 分离，避免双写互覆）和 paper_output/format_check_report.json"
            "（json 保持原名，workflow_guard.py 按该名消费）。"
            "退出码：0=PASS，1=FAIL。"
        ),
    )
    parser.add_argument(
        "--source",
        default=None,
        metavar="PATH",
        help=(
            "指定论文源文件（Markdown）。显式指定时正文文本改用该文件"
            "（跳过 docx 优先逻辑；Word 结构/字体检查仍针对 final_paper.docx）。"
            "默认不指定，保持原有路径逻辑。"
        ),
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    configure_utf8_stdio()
    args = build_arg_parser().parse_args(argv)
    source_override = Path(args.source) if args.source else None
    report = evaluate(source_override)
    write_reports(report)
    print(f"正式论文格式检查报告：{rel(REPORT_MD)}")
    if report["status"] == "PASS":
        print("✅ 正式论文格式门禁通过。")
        return 0
    print("⚠️ 正式论文格式门禁未通过。")
    for failure in report["failures"][:12]:
        print(f" - {failure}")
    if len(report["failures"]) > 12:
        print(f" - 其余 {len(report['failures']) - 12} 项见报告。")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
